"""ATS-safe .docx rendering — Phase 5.

Renders a CV and cover letter with python-docx. ATS rules enforced here:
  * single column, NO tables, NO text boxes, NO headers/footers
  * standard section headings (SUMMARY, EXPERIENCE, EDUCATION, ...)
  * CV bullets are taken VERBATIM from profile.yaml — only the summary (and the
    cover letter) come from the LLM, so nothing is fabricated.

We do NOT claim "100% ATS pass" — verify.py re-parses the output and asserts the
key fields are recoverable (the realistic guarantee).
"""

from __future__ import annotations

import os
from typing import Any

from docx import Document

from agent.tailor.generate import TailoredContent


def _fmt_dates(start: str | None, end: str | None) -> str:
    return f"{start or ''} – {end or 'Present'}".strip()


def _contact_line(personal: dict[str, Any]) -> str:
    bits = [
        personal.get("email"),
        personal.get("phone"),
        personal.get("location"),
        personal.get("linkedin_url"),  # display only — never automated
    ]
    return "  |  ".join(b for b in bits if b)


def build_cv(
    profile: dict[str, Any],
    job: dict[str, Any],
    tailored: TailoredContent,
    out_path: str,
) -> str:
    """Render an ATS-safe CV to out_path; return out_path."""
    personal = profile.get("personal", {})
    cv = profile.get("cv_content", {})
    market_tag = job.get("market_tag") or "eu"
    framing = (profile.get("market_framings", {}) or {}).get(market_tag, {}) \
        or (profile.get("market_framings", {}) or {}).get("eu", {})

    doc = Document()

    # Name + headline + contact (plain paragraphs, no header/footer)
    doc.add_heading(personal.get("name", ""), level=0)
    if framing.get("cv_headline"):
        doc.add_paragraph(framing["cv_headline"])
    doc.add_paragraph(_contact_line(personal))

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(tailored.summary.strip())

    if tailored.emphasis_keywords:
        doc.add_heading("Key Skills", level=1)
        doc.add_paragraph(", ".join(tailored.emphasis_keywords))

    doc.add_heading("Experience", level=1)
    for e in cv.get("experience", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{e.get('title')} — {e.get('company')}")
        run.bold = True
        meta = f"{e.get('location', '')}  |  {_fmt_dates(e.get('start'), e.get('end'))}"
        doc.add_paragraph(meta.strip())
        for bullet in e.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")

    if cv.get("education"):
        doc.add_heading("Education", level=1)
        for ed in cv["education"]:
            doc.add_paragraph(
                f"{ed.get('degree')} — {ed.get('institution')} ({ed.get('year')})"
            )

    if cv.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for c in cv["certifications"]:
            doc.add_paragraph(f"{c.get('name')} — {c.get('issuer')} ({c.get('year')})")

    if cv.get("skills"):
        doc.add_heading("Skills", level=1)
        for group, items in cv["skills"].items():
            label = group.replace("_", " ").title()
            doc.add_paragraph(f"{label}: {', '.join(items)}")

    if cv.get("languages"):
        doc.add_heading("Languages", level=1)
        doc.add_paragraph(
            ", ".join(f"{l.get('language')} ({l.get('level')})" for l in cv["languages"])
        )

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path


def build_cover_letter(
    profile: dict[str, Any],
    job: dict[str, Any],
    tailored: TailoredContent,
    out_path: str,
) -> str:
    """Render the cover letter to out_path; return out_path."""
    personal = profile.get("personal", {})

    doc = Document()
    doc.add_paragraph(personal.get("name", ""))
    doc.add_paragraph(_contact_line(personal))
    doc.add_paragraph("")
    doc.add_paragraph(f"Re: {job.get('title')} — {job.get('company')}")
    doc.add_paragraph("")

    for para in tailored.cover_letter.strip().split("\n\n"):
        doc.add_paragraph(para.strip())

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path
